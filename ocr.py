import os
import base64
import json
import fitz  # PyMuPDF
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import RGBColor, Inches
from openai import OpenAI


def encode_image(img):
    buffered = BytesIO()
    img.save(buffered, format="JPEG", quality=85)
    return base64.b64encode(buffered.getvalue()).decode('utf-8')


def count_pages(filepath):
    """Count pages without loading any page images into memory."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        with fitz.open(filepath) as doc:
            return len(doc)
    return 1


def extract_pages(filepath):
    """Yield page images one at a time so only a single page is ever in memory."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        doc = fitz.open(filepath)
        try:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # 150 dpi is plenty for GPT-4o (it downsizes to ~2048px anyway)
                # and uses a quarter of the memory of 300 dpi.
                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                pix = None  # release pixmap buffer promptly
                img.thumbnail((2000, 2000))
                yield img
        finally:
            doc.close()
    elif ext in ['.png', '.jpg', '.jpeg']:
        img = Image.open(filepath).convert('RGB')
        img.thumbnail((2000, 2000))
        yield img


def process_exam_paper(filepath, output_path, api_key, include_images, progress_callback):
    # Initialize the client dynamically with the user's provided key
    client = OpenAI(api_key=api_key)

    total_pages = count_pages(filepath)
    doc = Document()

    for i, img in enumerate(extract_pages(filepath)):
        progress_callback(int((i / total_pages) * 90))  # Allocate 90% of progress to API calls

        # Optionally embed the scanned page into the Word doc
        if include_images:
            img_path = f"{filepath}_temp_page_{i}.jpg"
            img_copy = img.copy()
            img_copy.thumbnail((600, 600))
            img_copy.save(img_path)
            doc.add_picture(img_path, width=Inches(6.0))
            os.remove(img_path)

        base64_image = encode_image(img)
        img = None  # free the page image before waiting on the API

        # Send to OpenAI
        response = client.chat.completions.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-4o"),
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": "You are a transcriber converting scanned exam papers. Return a JSON object with a 'blocks' array. Each block must have 'type' ('printed' for standard question text, 'handwritten' for student answers/working) and 'text'. Include crossed out text wrapped in ~~. Describe sketches in brackets."
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Transcribe this exam page."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }
            ],
            max_tokens=4000
        )

        try:
            content = json.loads(response.choices[0].message.content)
            blocks = content.get('blocks', [])

            # Format the output document
            for block in blocks:
                p = doc.add_paragraph()
                run = p.add_run(block.get('text', ''))

                if block.get('type') == 'handwritten':
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for handwriting
                    run.italic = True
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black for printed text

        except json.JSONDecodeError:
            doc.add_paragraph("[Error parsing transcription JSON structure for this page]")

        if i < total_pages - 1:
            doc.add_page_break()

    progress_callback(95)  # Finalizing document
    doc.save(output_path)
    progress_callback(100)  # Complete
